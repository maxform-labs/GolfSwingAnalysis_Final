#!/usr/bin/env python3
"""
FINAL_IMPROVEMENT_REPORT.md를 PDF와 DOCX로 변환하는 스크립트
한글 폰트 지원
"""

import os
import sys
import re
from pathlib import Path

def convert_to_pdf():
    """Markdown을 PDF로 변환 (ReportLab 사용)"""
    try:
        import markdown
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, 
                                        PageBreak, Table, TableStyle, KeepTogether,
                                        Preformatted)
        from reportlab.lib import colors
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.colors import HexColor
        from reportlab.platypus.flowables import HRFlowable
    except ImportError as e:
        print(f"ERROR: Required library not found: {e}")
        print("Please install: pip install reportlab markdown")
        return False
    
    md_file = "FINAL_IMPROVEMENT_REPORT.md"
    pdf_file = "FINAL_IMPROVEMENT_REPORT.pdf"
    
    if not os.path.exists(md_file):
        print(f"❌ 파일을 찾을 수 없습니다: {md_file}")
        return False
    
    # 한글 폰트 설정
    def setup_korean_fonts():
        """한글 폰트 설정"""
        try:
            # Windows 기본 한글 폰트
            pdfmetrics.registerFont(TTFont('MalgunGothic', 'c:/Windows/Fonts/malgun.ttf'))
            pdfmetrics.registerFont(TTFont('MalgunGothic-Bold', 'c:/Windows/Fonts/malgunbd.ttf'))
            return True
        except:
            try:
                pdfmetrics.registerFont(TTFont('MalgunGothic', 'c:/Windows/Fonts/gulim.ttc'))
                pdfmetrics.registerFont(TTFont('MalgunGothic-Bold', 'c:/Windows/Fonts/gulim.ttc'))
                return True
            except:
                print("WARNING: Korean font not found. Using default font.")
                return False
    
    has_korean_font = setup_korean_fonts()
    font_name = 'MalgunGothic' if has_korean_font else 'Helvetica'
    font_bold = 'MalgunGothic-Bold' if has_korean_font else 'Helvetica-Bold'
    
    # 스타일 생성
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name='CustomTitle',
        parent=styles['Title'],
        fontName=font_bold,
        fontSize=24,
        textColor=HexColor('#2c3e50'),
        spaceAfter=30,
        alignment=TA_CENTER
    ))
    
    styles.add(ParagraphStyle(
        name='SectionTitle',
        parent=styles['Heading1'],
        fontName=font_bold,
        fontSize=18,
        textColor=HexColor('#34495e'),
        spaceAfter=12,
        spaceBefore=20
    ))
    
    styles.add(ParagraphStyle(
        name='SubsectionTitle',
        parent=styles['Heading2'],
        fontName=font_bold,
        fontSize=16,
        textColor=HexColor('#7f8c8d'),
        spaceAfter=10,
        spaceBefore=15
    ))
    
    styles.add(ParagraphStyle(
        name='CustomBody',
        parent=styles['BodyText'],
        fontName=font_name,
        fontSize=11,
        leading=16,
        alignment=TA_JUSTIFY,
        spaceAfter=6
    ))
    
    # 코드 블록 스타일 - 한글 폰트 지원
    code_font_name = font_name if has_korean_font else 'Courier'
    styles.add(ParagraphStyle(
        name='CodeBlock',
        parent=styles['Code'],
        fontName=code_font_name,  # 한글 폰트 사용
        fontSize=9,
        leading=11,
        leftIndent=20,
        rightIndent=20,
        spaceBefore=6,
        spaceAfter=6,
        backColor=HexColor('#f8f8f8')
    ))
    
    # Markdown 파일 읽기
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Markdown을 HTML로 변환
    html = markdown.markdown(md_content, extensions=['tables', 'codehilite', 'fenced_code'])
    
    # PDF 문서 생성
    doc = SimpleDocTemplate(pdf_file, pagesize=A4,
                           rightMargin=2*cm, leftMargin=2*cm,
                           topMargin=2*cm, bottomMargin=2*cm)
    
    # 스토리 (내용) 생성
    story = []
    
    # HTML을 파싱하여 ReportLab 요소로 변환
    from html.parser import HTMLParser
    import re
    
    # 간단한 HTML 파서를 사용하여 Markdown의 HTML 출력을 처리
    # 실제로는 markdown의 HTML 출력을 직접 파싱하는 것이 더 복잡하므로
    # 마크다운을 직접 파싱하는 방식으로 변경
    
    # 강조 처리 함수
    def process_markdown_text(text):
        """Markdown 텍스트를 ReportLab HTML로 변환"""
        # HTML 태그 제거
        text = re.sub(r'<[^>]+>', '', text)
        # **text** -> <b>text</b>
        text = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', text)
        # 이모지 변환
        text = text.replace('✅', '✓').replace('❌', '✗').replace('⚠️', '!')
        return text
    
    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # 코드 블록 처리
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_block_lines = []
            else:
                in_code_block = False
                if code_block_lines:
                    # 한글을 포함한 코드 블록은 Paragraph로 처리
                    # 각 줄을 별도로 처리하여 한글 폰트 지원
                    for code_line in code_block_lines:
                        # HTML 특수문자 이스케이프
                        escaped_line = (code_line.replace('&', '&amp;')
                                       .replace('<', '&lt;')
                                       .replace('>', '&gt;'))
                        # Paragraph 사용 (한글 폰트 자동 지원)
                        p = Paragraph(escaped_line, styles['CodeBlock'])
                        story.append(p)
                    story.append(Spacer(1, 6))
                code_block_lines = []
            i += 1
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # 빈 줄
        if not line.strip():
            story.append(Spacer(1, 6))
            i += 1
            continue
        
        # 제목 처리
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            
            # HTML 태그 제거
            text = re.sub(r'<[^>]+>', '', text)
            
            if level == 1:
                story.append(Paragraph(text, styles['CustomTitle']))
            elif level == 2:
                story.append(Paragraph(text, styles['SectionTitle']))
            elif level == 3:
                story.append(Paragraph(text, styles['SubsectionTitle']))
            else:
                story.append(Paragraph(text, styles['CustomBody']))
            
            story.append(Spacer(1, 12))
            i += 1
            continue
        
        # 구분선
        if line.strip() == '---':
            story.append(HRFlowable(width="100%", thickness=1, lineCap='round', color=colors.grey))
            story.append(Spacer(1, 12))
            i += 1
            continue
        
        # 리스트 처리
        if line.strip().startswith('-') or line.strip().startswith('*'):
            text = line.strip()[1:].strip()
            text = process_markdown_text(text)
            p = Paragraph(f"• {text}", styles['CustomBody'])
            story.append(p)
            i += 1
            continue
        
        # 번호 리스트
        if re.match(r'^\d+\.', line.strip()):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            text = process_markdown_text(text)
            p = Paragraph(text, styles['CustomBody'])
            story.append(p)
            i += 1
            continue
        
        # 일반 텍스트
        text = process_markdown_text(line)
        
        if text.strip():
            p = Paragraph(text, styles['CustomBody'])
            story.append(p)
        
        i += 1
    
    try:
        # PDF 생성
        doc.build(story)
        print(f"SUCCESS: PDF conversion completed: {pdf_file}")
        return True
    except Exception as e:
        print(f"ERROR: PDF conversion failed: {e}")
        import traceback
        traceback.print_exc()
        return False


def convert_to_docx():
    """Markdown을 DOCX로 변환 (python-docx 사용)"""
    try:
        import markdown
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        print("ERROR: Required library not found.")
        print("Please install: pip install markdown python-docx")
        return False
    
    md_file = "FINAL_IMPROVEMENT_REPORT.md"
    docx_file = "FINAL_IMPROVEMENT_REPORT.docx"
    
    if not os.path.exists(md_file):
        print(f"❌ 파일을 찾을 수 없습니다: {md_file}")
        return False
    
    # Markdown 파일 읽기
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 새 문서 생성
    doc = Document()
    
    # 한글 폰트 설정
    def set_font(run, font_name='Malgun Gothic', size=None):
        """한글 폰트 설정"""
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        if size:
            run.font.size = Pt(size)
    
    # 스타일 설정
    styles = doc.styles
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Malgun Gothic'
    normal_font._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    normal_font.size = Pt(11)
    
    # Markdown 파싱 및 변환
    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []
    code_language = ''
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # 코드 블록 처리
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_language = line[3:].strip()
                code_block_lines = []
            else:
                in_code_block = False
                # 코드 블록 추가
                if code_block_lines:
                    # 각 줄을 별도 문단으로 추가 (한글 지원)
                    for code_line in code_block_lines:
                        p = doc.add_paragraph(style='No Spacing')
                        # 한글 폰트 설정
                        run = p.add_run(code_line)
                        run.font.name = 'Courier New'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')  # 한글 폰트
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        # 코드 블록 스타일 적용
                        p.paragraph_format.left_indent = Inches(0.5)
                        p.paragraph_format.space_before = Pt(6)
                        p.paragraph_format.space_after = Pt(6)
                code_block_lines = []
            i += 1
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # 빈 줄
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue
        
        # 제목 처리
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            
            if level == 1:
                p = doc.add_heading(text, level=1)
            elif level == 2:
                p = doc.add_heading(text, level=2)
            elif level == 3:
                p = doc.add_heading(text, level=3)
            elif level == 4:
                p = doc.add_heading(text, level=4)
            else:
                p = doc.add_paragraph(text)
            
            # 폰트 설정
            for run in p.runs:
                set_font(run, 'Malgun Gothic')
            
            i += 1
            continue
        
        # 리스트 처리
        if line.strip().startswith('-') or line.strip().startswith('*'):
            text = line.strip()[1:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
            for run in p.runs:
                set_font(run, 'Malgun Gothic')
            i += 1
            continue
        
        # 번호 리스트
        if re.match(r'^\d+\.', line.strip()):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            p = doc.add_paragraph(text, style='List Number')
            for run in p.runs:
                set_font(run, 'Malgun Gothic')
            i += 1
            continue
        
        # 테이블 처리 (간단한 형태만)
        if '|' in line and line.count('|') >= 2:
            # 테이블 헤더/구분선/데이터 처리
            table_lines = [line]
            j = i + 1
            while j < len(lines) and '|' in lines[j] and lines[j].count('|') >= 2:
                if not lines[j].strip().startswith('|--'):
                    table_lines.append(lines[j])
                j += 1
            
            if len(table_lines) > 1:
                # 테이블 생성
                rows = []
                for tl in table_lines:
                    if tl.strip().startswith('|--'):
                        continue
                    cells = [c.strip() for c in tl.split('|')[1:-1]]
                    rows.append(cells)
                
                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    for row_idx, row in enumerate(rows):
                        for col_idx, cell_text in enumerate(row):
                            cell = table.rows[row_idx].cells[col_idx]
                            p = cell.paragraphs[0]
                            run = p.add_run(cell_text)
                            set_font(run, 'Malgun Gothic', 10)
                            if row_idx == 0:  # 헤더
                                run.font.bold = True
                
                i = j
                continue
        
        # 일반 텍스트
        p = doc.add_paragraph()
        
        # 강조 처리
        text = line
        # **bold** 처리
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                set_font(run, 'Malgun Gothic', 11)
                run.font.bold = True
            else:
                run = p.add_run(part)
                set_font(run, 'Malgun Gothic', 11)
        
        i += 1
    
    # 문서 저장
    try:
        doc.save(docx_file)
        print(f"SUCCESS: DOCX conversion completed: {docx_file}")
        return True
    except Exception as e:
        print(f"ERROR: DOCX conversion failed: {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    """메인 함수"""
    print("=" * 60)
    print("Converting FINAL_IMPROVEMENT_REPORT.md")
    print("=" * 60)
    
    # PDF 변환
    print("\n[1/2] Converting to PDF...")
    pdf_success = convert_to_pdf()
    
    # DOCX 변환
    print("\n[2/2] Converting to DOCX...")
    docx_success = convert_to_docx()
    
    print("\n" + "=" * 60)
    print("Conversion completed")
    print("=" * 60)
    if pdf_success:
        print(f"SUCCESS: PDF: FINAL_IMPROVEMENT_REPORT.pdf")
    if docx_success:
        print(f"SUCCESS: DOCX: FINAL_IMPROVEMENT_REPORT.docx")
    
    if not pdf_success and not docx_success:
        print("\nERROR: All conversions failed.")
        print("Please install required libraries:")
        print("  pip install reportlab markdown python-docx")
        return 1
    
    return 0


if __name__ == "__main__":
    sys.exit(main())

